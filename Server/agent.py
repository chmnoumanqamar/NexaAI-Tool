import os
import re
import json
import base64
import zipfile
import requests
import pandas as pd
from dotenv import load_dotenv
from reportlab.pdfgen import canvas
from docx import Document
from langchain_core.tools import tool
from duckduckgo_search import DDGS
from langchain_community.tools import DuckDuckGoSearchRun
from langchain_groq import ChatGroq
from langchain.agents import create_agent
from groq import Groq
import wikipedia

load_dotenv()

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "secure_uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

_CURRENT_WORKSPACE_ID: str | None = None

def set_current_workspace(ws_id: str | None):
    global _CURRENT_WORKSPACE_ID
    _CURRENT_WORKSPACE_ID = ws_id

def get_upload_dir() -> str:
    if _CURRENT_WORKSPACE_ID:
        ws_path = os.path.join(UPLOAD_DIR, _CURRENT_WORKSPACE_ID)
        os.makedirs(ws_path, exist_ok=True)
        return ws_path
    return UPLOAD_DIR

def get_download_url(filename: str) -> str:
    safe_fn = os.path.basename(filename)
    if _CURRENT_WORKSPACE_ID:
        return f"http://localhost:8000/api/download/{safe_fn}?workspace_id={_CURRENT_WORKSPACE_ID}"
    return f"http://localhost:8000/api/download/{safe_fn}"

def find_file_in_workspace(filename: str) -> str | None:
    safe_fn = os.path.basename(filename)
    if _CURRENT_WORKSPACE_ID:
        ws_fp = os.path.join(UPLOAD_DIR, _CURRENT_WORKSPACE_ID, safe_fn)
        if os.path.exists(ws_fp):
            return ws_fp
    root_fp = os.path.join(UPLOAD_DIR, safe_fn)
    if os.path.exists(root_fp):
        return root_fp
    for root, _, files in os.walk(UPLOAD_DIR):
        if safe_fn in files:
            return os.path.join(root, safe_fn)
    return None

groq_api_key = os.getenv("GROQ_API_KEY", "")
groq_client = Groq(api_key=groq_api_key) if groq_api_key else None

WMO_CODE_MAP = {
    0: "Clear sky ☀️",
    1: "Mainly clear 🌤️",
    2: "Partly cloudy ⛅",
    3: "Overcast ☁️",
    45: "Foggy 🌫️",
    48: "Depositing rime fog 🌫️",
    51: "Light drizzle 🌦️",
    53: "Moderate drizzle 🌦️",
    55: "Dense drizzle 🌧️",
    61: "Slight rain 🌧️",
    63: "Moderate rain 🌧️",
    65: "Heavy rain 🌧️",
    71: "Slight snow 🌨️",
    73: "Moderate snow 🌨️",
    75: "Heavy snow ❄️",
    80: "Slight rain showers 🌦️",
    81: "Moderate rain showers 🌧️",
    82: "Violent rain showers ⛈️",
    95: "Thunderstorm ⛈️",
    96: "Thunderstorm with slight hail ⛈️",
    99: "Thunderstorm with heavy hail ⛈️"
}

# ============================================================
# TOOL 1: Real-time Live Weather Tool (Open-Meteo API + wttr.in Fallback)
# ============================================================
def _clean_location_query(raw_loc: str) -> str:
    """Cleans user location strings in English, Urdu, Roman Urdu, Hindi, Spanish, French, etc."""
    text = raw_loc.strip()
    stopwords = [
        "weather in", "weather of", "weather at", "weather for", "weather",
        "mausam", "ka mausam", "ki weather", "current", "live", "today",
        "temperature in", "temperature of", "temperature", "forecast for", "forecast",
        "city", "batao", "bata do", "bataiye", "kaisa hai", "kya hai", "update", "now",
        "climat", "météo", "tiempo", "wetter", "ka hal", "halat", "check"
    ]
    for w in stopwords:
        text = re.sub(rf"\b{re.escape(w)}\b", "", text, flags=re.IGNORECASE)
    text = text.strip(" ,.-/?!")
    if "," in text:
        parts = [p.strip() for p in text.split(",") if p.strip()]
        if parts:
            text = parts[0]
    return text.strip() or raw_loc.strip()

@tool
def get_live_weather(location: str) -> str:
    """Fetches accurate real-time live weather data (current temperature in Celsius, feels-like temperature, condition, humidity %, wind speed, precipitation, today's high/low) for any city, town, or location in the world.
    MUST ALWAYS be used when user asks about current weather, temperature, rain, forecast, or climate for any city (e.g. Vehari, Lahore, Karachi, Islamabad, London, New York, etc.)."""
    cleaned = _clean_location_query(location)
    query_candidates = [cleaned, location.strip()] if cleaned != location.strip() else [cleaned]

    # Tier 1: Try Open-Meteo with geocoding
    for candidate in query_candidates:
        try:
            geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={requests.utils.quote(candidate)}&count=1&language=en&format=json"
            geo_res = requests.get(geo_url, timeout=6).json()
            results = geo_res.get("results")
            if results:
                place = results[0]
                name = place.get("name")
                country = place.get("country", "")
                admin1 = place.get("admin1", "")
                lat = place["latitude"]
                lon = place["longitude"]
                
                weather_url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current=temperature_2m,relative_humidity_2m,apparent_temperature,is_day,precipitation,weather_code,wind_speed_10m&daily=temperature_2m_max,temperature_2m_min,precipitation_probability_max&timezone=auto"
                w_res = requests.get(weather_url, timeout=6).json()
                
                curr = w_res.get("current", {})
                daily = w_res.get("daily", {})
                
                temp = curr.get("temperature_2m")
                feels = curr.get("apparent_temperature")
                humidity = curr.get("relative_humidity_2m")
                wind = curr.get("wind_speed_10m")
                precip = curr.get("precipitation", 0)
                code = curr.get("weather_code", 0)
                condition = WMO_CODE_MAP.get(code, "Clear")
                
                loc_str = f"{name}"
                if admin1 and admin1 != name:
                    loc_str += f", {admin1}"
                if country:
                    loc_str += f", {country}"
                    
                max_t = daily.get("temperature_2m_max", [temp])[0] if daily.get("temperature_2m_max") else temp
                min_t = daily.get("temperature_2m_min", [temp])[0] if daily.get("temperature_2m_min") else temp
                
                return (
                    f"📍 Real-time Weather Report for {loc_str}:\n"
                    f"- Condition: {condition}\n"
                    f"- Current Temperature: {temp}°C\n"
                    f"- Feels Like: {feels}°C\n"
                    f"- Today High / Low: {max_t}°C / {min_t}°C\n"
                    f"- Humidity: {humidity}%\n"
                    f"- Wind Speed: {wind} km/h\n"
                    f"- Precipitation: {precip} mm\n"
                    f"- Timezone: {w_res.get('timezone', 'Local')}"
                )
        except Exception:
            pass

    # Tier 2: Try wttr.in fallback
    try:
        wttr_url = f"https://wttr.in/{requests.utils.quote(cleaned)}?format=j1"
        w_data = requests.get(wttr_url, timeout=6).json()
        curr_cond = w_data.get("current_condition", [{}])[0]
        weather_list = w_data.get("weather", [{}])
        
        temp_c = curr_cond.get("temp_C")
        feels_c = curr_cond.get("FeelsLikeC", temp_c)
        humidity = curr_cond.get("humidity")
        wind_km = curr_cond.get("windspeedKmph")
        precip_mm = curr_cond.get("precipMM", "0.0")
        desc = curr_cond.get("weatherDesc", [{}])[0].get("value", "Clear")
        
        max_t = weather_list[0].get("maxtempC", temp_c) if weather_list else temp_c
        min_t = weather_list[0].get("mintempC", temp_c) if weather_list else temp_c
        
        area = w_data.get("nearest_area", [{}])[0]
        area_name = area.get("areaName", [{}])[0].get("value", cleaned.title())
        country_name = area.get("country", [{}])[0].get("value", "")
        loc_display = f"{area_name}, {country_name}" if country_name else area_name

        return (
            f"📍 Real-time Weather Report for {loc_display}:\n"
            f"- Condition: {desc} 🌤️\n"
            f"- Current Temperature: {temp_c}°C\n"
            f"- Feels Like: {feels_c}°C\n"
            f"- Today High / Low: {max_t}°C / {min_t}°C\n"
            f"- Humidity: {humidity}%\n"
            f"- Wind Speed: {wind_km} km/h\n"
            f"- Precipitation: {precip_mm} mm"
        )
    except Exception as e:
        return f"Could not retrieve live weather for '{location}': {str(e)}. Please check the city name."

# ============================================================
# TOOL 2: Real-time Internet Web Search
# ============================================================
@tool
def search_internet(query: str) -> str:
    """Searches the internet in real-time for latest news, facts, current events, live updates, or answers.
    Use this when user asks about current information, recent events, live scores, stock rates, or web data."""
    try:
        results = DDGS().text(query, max_results=5)
        if results:
            snippets = [f"• {r.get('title')}: {r.get('body')} ({r.get('href')})" for r in results]
            return "\n\n".join(snippets)
        return DuckDuckGoSearchRun().run(query)
    except Exception:
        try:
            return DuckDuckGoSearchRun().run(query)
        except Exception as e2:
            return f"Web search error: {str(e2)}"

# ============================================================
# TOOL 3: Wikipedia Knowledge Summary
# ============================================================
@tool
def get_wikipedia_summary(query: str) -> str:
    """Fetches a factual encyclopedia summary of a topic, person, place, or concept from Wikipedia.
    Use this for definitions, history, biographies, and factual overviews."""
    try:
        wikipedia.set_user_agent("NexaAI-Assistant/2.0 (contact@nexaai.com)")
        wikipedia.set_lang("en")
        return wikipedia.summary(query, sentences=4)
    except Exception:
        try:
            clean_q = requests.utils.quote(query.strip())
            r = requests.get(
                f"https://en.wikipedia.org/api/rest_v1/page/summary/{clean_q}",
                headers={"User-Agent": "NexaAI-Assistant/2.0 (contact@nexaai.com)"},
                timeout=6
            )
            if r.status_code == 200:
                data = r.json()
                extract = data.get("extract")
                if extract:
                    return extract
            return f"No direct Wikipedia page found for '{query}'."
        except Exception as e2:
            return f"Wikipedia search notice: {str(e2)}"

# ============================================================
# TOOL 4: Generate Excel (.xlsx) File
# ============================================================
@tool
def generate_excel(data_json: str, filename: str) -> str:
    """Converts structured JSON data into a Microsoft Excel (.xlsx) spreadsheet.
    Input data_json must be a valid JSON string (list of dicts, e.g. '[{"Item": "A", "Price": 100}]' or dict with lists).
    filename should end with .xlsx (e.g. 'budget_report.xlsx').
    CRITICAL: ONLY invoke when the user EXPLICITLY asks to create, export, or download an Excel file."""
    try:
        fn = filename if filename.lower().endswith(".xlsx") else f"{filename}.xlsx"
        parsed = json.loads(data_json) if isinstance(data_json, str) else data_json
        if isinstance(parsed, list):
            df = pd.DataFrame(parsed)
        elif isinstance(parsed, dict):
            df = pd.DataFrame(parsed)
        else:
            df = pd.DataFrame([{"Content": str(parsed)}])
        out_path = os.path.join(get_upload_dir(), fn)
        with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Sheet1")
        dl = get_download_url(fn)
        return f"✅ Excel spreadsheet '{fn}' generated successfully! Download Link: {dl}"
    except Exception as e:
        return f"Error generating Excel file: {str(e)}"

# ============================================================
# TOOL 5: Generate CSV (.csv) File
# ============================================================
@tool
def generate_csv(data_json: str, filename: str) -> str:
    """Converts structured JSON data into a CSV (.csv) data file.
    Input data_json must be a valid JSON string (list of dicts).
    filename should end with .csv (e.g. 'dataset.csv').
    CRITICAL: ONLY invoke when the user EXPLICITLY asks to create, export, or download a CSV file."""
    try:
        fn = filename if filename.lower().endswith(".csv") else f"{filename}.csv"
        parsed = json.loads(data_json) if isinstance(data_json, str) else data_json
        df = pd.DataFrame(parsed)
        out_path = os.path.join(get_upload_dir(), fn)
        df.to_csv(out_path, index=False, encoding="utf-8")
        dl = get_download_url(fn)
        return f"✅ CSV file '{fn}' generated successfully! Download Link: {dl}"
    except Exception as e:
        return f"Error generating CSV file: {str(e)}"

# ============================================================
# TOOL 6: Generate PDF (.pdf) File
# ============================================================
@tool
def generate_pdf(text: str, filename: str) -> str:
    """Generates a styled, professional PDF document from given text content.
    filename should end with .pdf (e.g. 'weather_report.pdf' or 'summary.pdf').
    CRITICAL: ONLY invoke when the user EXPLICITLY asks to create, export, or download a PDF document."""
    try:
        fn = filename if filename.lower().endswith(".pdf") else f"{filename}.pdf"
        out_path = os.path.join(get_upload_dir(), fn)
        pdf = canvas.Canvas(out_path)
        pdf.setTitle(fn.replace(".pdf", "").replace("_", " ").title())
        
        # Header bar
        pdf.setFillColorRGB(0.12, 0.28, 0.18)  # Sage forest green
        pdf.rect(0, 792, 612, 50, fill=1, stroke=0)
        pdf.setFillColorRGB(1, 1, 1)
        pdf.setFont("Helvetica-Bold", 14)
        doc_title = fn.replace(".pdf", "").replace("_", " ").title()
        pdf.drawString(40, 808, f"NexaAI Document — {doc_title}")

        # Body
        pdf.setFillColorRGB(0.1, 0.15, 0.12)
        pdf.setFont("Helvetica", 10.5)
        y = 750
        margin_left = 45
        line_height = 16
        max_chars = 82

        lines = text.split("\n")
        for raw_line in lines:
            trimmed = raw_line.strip()
            if not trimmed:
                y -= 10
                if y < 60:
                    pdf.showPage()
                    y = 750
                    pdf.setFont("Helvetica", 10.5)
                continue

            # Heading detection
            if trimmed.startswith("# ") or trimmed.startswith("## ") or (trimmed.isupper() and len(trimmed) < 40):
                heading_txt = trimmed.lstrip("#").strip()
                y -= 8
                pdf.setFont("Helvetica-Bold", 12.5)
                pdf.setFillColorRGB(0.12, 0.35, 0.22)
                pdf.drawString(margin_left, y, heading_txt)
                y -= (line_height + 4)
                pdf.setFont("Helvetica", 10.5)
                pdf.setFillColorRGB(0.1, 0.15, 0.12)
            else:
                # Wrap long text
                words = trimmed.split(" ")
                current_line = ""
                for w in words:
                    test_line = f"{current_line} {w}".strip()
                    if len(test_line) > max_chars:
                        if y < 60:
                            pdf.showPage()
                            y = 750
                            pdf.setFont("Helvetica", 10.5)
                        pdf.drawString(margin_left, y, current_line)
                        y -= line_height
                        current_line = w
                    else:
                        current_line = test_line
                    if current_line:
                        if y < 60:
                            pdf.showPage()
                            y = 750
                            pdf.setFont("Helvetica", 10.5)
                        pdf.drawString(margin_left, y, current_line)
                        y -= line_height

        pdf.save()
        dl = get_download_url(fn)
        return f"✅ PDF document '{fn}' generated successfully! Download Link: {dl}"
    except Exception as e:
        return f"Error generating PDF document: {str(e)}"

# ============================================================
# TOOL 7: Generate Word (.docx) Document
# ============================================================
@tool
def generate_word(text: str, filename: str) -> str:
    """Generates a Microsoft Word (.docx) document with clean headings, bullets, and paragraphs.
    filename should end with .docx (e.g. 'project_report.docx').
    CRITICAL: ONLY invoke when the user EXPLICITLY asks to create, export, or download a Word document (.docx)."""
    try:
        fn = filename if filename.lower().endswith(".docx") else f"{filename}.docx"
        out_path = os.path.join(get_upload_dir(), fn)
        doc = Document()
        doc_title = fn.replace(".docx", "").replace("_", " ").title()
        doc.add_heading(doc_title, level=0)

        for paragraph in text.split("\n\n"):
            p_text = paragraph.strip()
            if not p_text:
                continue
            if p_text.startswith("# ") or p_text.startswith("## "):
                clean_h = p_text.lstrip("#").strip()
                doc.add_heading(clean_h, level=1)
            elif p_text.startswith("### "):
                clean_h = p_text.lstrip("#").strip()
                doc.add_heading(clean_h, level=2)
            elif p_text.startswith("• ") or p_text.startswith("- "):
                for bullet_line in p_text.split("\n"):
                    b_clean = bullet_line.lstrip("•-* ").strip()
                    if b_clean:
                        doc.add_paragraph(b_clean, style="List Bullet")
            else:
                doc.add_paragraph(p_text)

        doc.save(out_path)
        dl = get_download_url(fn)
        return f"✅ Word document '{fn}' generated successfully! Download Link: {dl}"
    except Exception as e:
        return f"Error generating Word document: {str(e)}"

# ============================================================
# TOOL 8: Zip Multiple Files (.zip)
# ============================================================
@tool
def zip_files(filenames_json: str, zip_filename: str) -> str:
    """Bundles multiple files from secure uploads into a single ZIP archive.
    filenames_json must be a JSON list of filenames, e.g. '["data.xlsx", "report.pdf"]'.
    zip_filename should end with .zip (e.g. 'project_files.zip').
    CRITICAL: ONLY invoke when the user EXPLICITLY asks to zip or bundle files."""
    try:
        fn = zip_filename if zip_filename.lower().endswith(".zip") else f"{zip_filename}.zip"
        zip_path = os.path.join(get_upload_dir(), fn)
        filenames = json.loads(filenames_json) if isinstance(filenames_json, str) else filenames_json
        found = []
        missing = []
        with zipfile.ZipFile(zip_path, "w") as zf:
            for name in filenames:
                fp = find_file_in_workspace(name)
                if fp and os.path.exists(fp):
                    zf.write(fp, arcname=name)
                    found.append(name)
                else:
                    missing.append(name)
        dl = get_download_url(fn)
        msg = f"✅ ZIP archive '{fn}' created with {len(found)} file(s)."
        if missing:
            msg += f" (Missing: {', '.join(missing)})."
        msg += f" Download Link: {dl}"
        return msg
    except Exception as e:
        return f"Error creating ZIP archive: {str(e)}"

# ============================================================
# TOOL 9: Analyze Uploaded Image (Vision)
# ============================================================
@tool
def analyze_uploaded_image(image_filename: str, question: str = "Describe this image in detail and extract all key data.") -> str:
    """Analyzes an uploaded image and answers questions about it using AI vision.
    image_filename must be a filename present in secure uploads.
    Use this when user uploads an image and asks to read, analyze, or transcribe it."""
    if not groq_client:
        return "Groq client not configured with API key."
    try:
        image_path = find_file_in_workspace(image_filename)
        if not image_path or not os.path.exists(image_path):
            return f"Image '{image_filename}' not found in uploads directory. Please ensure it was uploaded."
        with open(image_path, "rb") as f:
            b64_img = base64.b64encode(f.read()).decode("utf-8")
        ext = image_filename.lower().split(".")[-1]
        mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "webp": "image/webp", "gif": "image/gif"}.get(ext, "image/jpeg")

        res = groq_client.chat.completions.create(
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": question or "Extract all information and describe this image."},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64_img}"}}
                ]
            }],
            model="qwen/qwen3.6-27b",
        )
        return res.choices[0].message.content
    except Exception as e:
        return f"Image vision error: {str(e)}"

# ============================================================
# TOOL 10: List Workspace Files
# ============================================================
@tool
def list_workspace_files() -> str:
    """Lists all files currently available in the workspace (uploaded files, generated PDFs, Excel sheets, Word docs, etc.).
    Use this when the user asks what files exist or refers to previous files."""
    try:
        target_dir = get_upload_dir()
        files = [f for f in os.listdir(target_dir) if not f.startswith(".") and f.lower() != ".gitkeep"]
        if not files:
            return "No files currently uploaded or generated in this workspace."
        return f"Files available in workspace: {', '.join(files)}"
    except Exception as e:
        return f"Error listing workspace files: {str(e)}"

# ============================================================
# TOOL 11: Read Workspace File Content
# ============================================================
@tool
def read_workspace_file(filename: str) -> str:
    """Reads the text content of a file (e.g. .txt, .csv, .py, .js, .json, .md, .html) from workspace storage.
    Use this when user asks to read, inspect, or summarize a text/code/csv file in the workspace."""
    try:
        safe_fn = os.path.basename(filename)
        path = find_file_in_workspace(safe_fn)
        if not path or not os.path.exists(path):
            return f"File '{safe_fn}' not found in workspace."
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read(6000)
        return f"Content of '{safe_fn}':\n{content}"
    except Exception as e:
        return f"Error reading file '{filename}': {str(e)}"

# ============================================================
# Agent Setup
# ============================================================
agent_tools = [
    get_live_weather,
    search_internet,
    get_wikipedia_summary,
    generate_excel,
    generate_csv,
    generate_pdf,
    generate_word,
    zip_files,
    analyze_uploaded_image,
    list_workspace_files,
    read_workspace_file
]

system_prompt = (
    "You are NexaAI, an advanced, highly intelligent multilingual AI assistant operating in a collaborative workspace.\n\n"
    "TOOL CAPABILITIES:\n"
    "1. `get_live_weather`: Fetches real-time temperature, condition, humidity, and forecast for any city or town worldwide.\n"
    "2. `search_internet`: Real-time DuckDuckGo web search for live news, scores, currency exchange, prices, and facts.\n"
    "3. `get_wikipedia_summary`: Factual encyclopedia overviews for concepts, history, and scientific topics.\n"
    "4. `generate_pdf`: Generates styled PDF (.pdf) documents.\n"
    "5. `generate_word`: Generates Microsoft Word (.docx) documents.\n"
    "6. `generate_excel`: Generates Microsoft Excel (.xlsx) spreadsheets.\n"
    "7. `generate_csv`: Generates CSV (.csv) data files.\n"
    "8. `zip_files`: Bundles files into a ZIP archive (.zip).\n"
    "9. `analyze_uploaded_image`: Vision tool to inspect user uploaded images.\n"
    "10. `list_workspace_files`: Inspects and lists all files currently uploaded or generated in the workspace.\n"
    "11. `read_workspace_file`: Reads text/code/data from files in the workspace.\n\n"
    "STRICT CORE RULES (READ CAREFULLY):\n"
    "- CONVERSATION CONTEXT & MEMORY: You have full access to previous messages in the current conversation. When the user asks follow-up questions (e.g. 'give me the third file', 'explain step 2', 'app.py', 'continue', 'make changes to it'), always refer to what was previously discussed, provide the requested code/information thoroughly, and never claim you lack context.\n"
    "- MULTI-LANGUAGE RESPONSIVENESS (MANDATORY): Always detect the language and script of the user's message (English, Roman Urdu / Roman Hindi, Urdu script اردو, Hindi हिन्दी, Arabic العربية, Spanish, French, German, Turkish, Chinese, etc.) and ALWAYS reply fluently, accurately, and naturally in that EXACT SAME language and script.\n"
    "  * If the user speaks in English, reply in pure, elegant English.\n"
    "  * If the user speaks in Roman Urdu / Roman Hindi (e.g. 'vehari ka mausam kaisa hai', 'kya haal hai', 'mujhe ek report chahiye'), reply naturally and clearly in Roman Urdu.\n"
    "  * If the user writes in Urdu script (اردو), reply in Urdu script.\n"
    "  * If the user writes in Arabic, Spanish, French, etc., reply in that specific language.\n"
    "  * Never mix languages unless explicitly requested by the user.\n"
    "- NEVER generate files (PDF, Word, Excel, CSV, ZIP) automatically for normal questions, searches, weather requests, summaries, code inquiries, or chit-chat. For all standard queries, reply directly in the chat with clean, well-formatted text or code blocks.\n"
    "- ONLY invoke file generation tools (`generate_pdf`, `generate_word`, `generate_excel`, `generate_csv`, `zip_files`) when the user EXPLICITLY commands you to make, create, export, or download a file (e.g. 'make a PDF', 'Word document banao', 'export as Excel', 'download as CSV', 'crear PDF', etc.).\n"
    "- When you generate a file upon user request, give a helpful summary of what was generated in the user's language and ALWAYS include the download link returned by the tool (http://localhost:8000/api/download/<filename>).\n"
    "- Format text cleanly using markdown bolding, bullet points, headers, and organized code blocks."
)

llm = ChatGroq(
    temperature=0,
    model_name="qwen/qwen3.6-27b",
    api_key=groq_api_key
)

agent = create_agent(llm, agent_tools, system_prompt=system_prompt)